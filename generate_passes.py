"""
Автоматическая генерация разовых пропусков и отправка в Telegram.
Запускается ежедневно в 21:00 МСК через GitHub Actions.
Генерирует два документа:
1. Официальное письмо в ГБУ «Мосприрода»
2. Перечень автомобилей (простая таблица)
"""

import json
import logging
import os
import sys
import urllib.request
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from pathlib import Path

import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from num_to_words import number_to_genitive

SPREADSHEET_ID = os.environ.get("SPREADSHEET_ID", "1ZGR8581dQu-rhvgnOBJ0AaIOqN6z-PKapZSa_k2qRkU")
SHEET_NAME = "Реестр автомобилей"
BOT_TOKEN = os.environ.get("BOT_TOKEN", "")
RECIPIENT_CHAT_ID = os.environ.get("RECIPIENT_CHAT_ID", "-5265944992")
TIMEZONE = ZoneInfo("Europe/Moscow")
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/drive.readonly",
]

logging.basicConfig(format="%(asctime)s - %(levelname)s - %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)


def get_cars_for_date(target_date):
    """Получает из Google Sheets список авто на указанную дату.
    При дублировании гос. номера на одну дату берётся последняя заявка."""
    creds_json = os.environ.get("GOOGLE_CREDENTIALS")
    if not creds_json:
        logger.error("GOOGLE_CREDENTIALS не задана")
        return []
    creds = Credentials.from_service_account_info(json.loads(creds_json), scopes=SCOPES)
    client = gspread.authorize(creds)
    sheet = client.open_by_key(SPREADSHEET_ID).worksheet(SHEET_NAME)
    records = sheet.get_all_records()

    # Словарь для дедупликации: ключ — гос. номер в верхнем регистре,
    # значение — данные авто. Последняя запись перезаписывает предыдущую.
    cars_by_number = {}
    for row in records:
        arrival = str(row.get("Дата прибытия", "")).strip()
        if arrival == target_date:
            brand = str(row.get("Марка", "")).strip()
            number = str(row.get("Гос. номер", "")).strip()
            if brand and number:
                cars_by_number[number.upper()] = {"car_brand": brand, "car_number": number}

    cars = list(cars_by_number.values())
    logger.info(f"Найдено {len(cars)} уникальных авто на {target_date}")
    return cars


def set_cell_border(cell):
    """Устанавливает границы ячейки таблицы."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge in ["top", "bottom", "left", "right"]:
        el = tcBorders.makeelement(
            qn(f"w:{edge}"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): "000000"},
        )
        tcBorders.append(el)


def generate_document(cars, target_date_str):
    """Генерирует Word-документ с официальным письмом."""
    target_date = datetime.strptime(target_date_str, "%d.%m.%Y")
    doc_date = datetime.now(TIMEZONE)
    doc_number = doc_date.strftime("%d-%m/%y")
    doc_date_fmt = doc_date.strftime("%d.%m.%Y")
    count = len(cars)
    count_words = number_to_genitive(count)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(f"{doc_date_fmt}. № {doc_number}")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("На №________ от________")
    run.font.size = Pt(12)

    doc.add_paragraph()

    for text in ["Генеральному директору", "ГБУ «Мосприрода»"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Адигамовой Ю.И.")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Уважаемая Юлия Ильдусовна!")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("В рамках ранее направленного обращения АНО СК «Олимпик» от 11.03.2026 № 11-03/26,")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "в соответствии с договором аренды объекта особо ценного движимого имущества "
        "№ ОЦДИ-17, заключенным между ГПБУ «Мосприрода» и Автономной некоммерческой "
        "Организацией Спортивный клуб «Олимпик» (ОГРН 1047796891570) 01.08.2018 г., "
        "в целях выполнения разовых обязательств, прошу Вас дать разовые разрешение на "
        "въезд транспортных средств на особо охраняемую зеленую территорию ландшафтный "
        "заказник «Теплый Стан», подведомственную ГАУК г. Москвы «Парки Москвы», для "
        "проезда к объектам, расположенным по адресу: г. Москва, ул. Островитянова, вл.10., "
        f"на {target_date_str} г., согласно прилагаемой схеме движения автотранспорта."
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run("Информация о марках транспортных средств и государственных регистрационных знаках:")
    run.font.size = Pt(12)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["№", "Номер автомобиля", "Марка автомобиля"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        set_cell_border(cell)

    for idx, car in enumerate(cars, 1):
        row = table.add_row()
        for i, val in enumerate([str(idx), car["car_number"], car["car_brand"]]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = "Arial"
            set_cell_border(cell)

    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(7)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        f"Также сообщаю, что единовременно на территории ландшафтного заказника "
        f"«Теплый Стан» по адресу : г. Москва ул. Островитянова, вл.10, заезд "
        f"автомобилей не будет превышать количества {count} ({count_words}) штук, "
        f"а также будут соблюдаться регламент"
    )
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "постановления Правительства Москвы от 14.09.2010 №795-ПП «Об утверждении "
        "Регламента подготовки и выдачи заявителем Департаментом природопользования "
        "и охраны окружающей среды города Москвы на въезд на особо охраняемые природные "
        "территории города Москвы», а именно:"
    )
    run.font.size = Pt(12)

    for bullet in [
        "въезд и передвижение транспортных средств по ООЗТ вне дорог общего пользования будет осуществляться по строго установленным маршрутам (технологическим картам), согласованным балансодержателем территории ГАУК г. Москвы «Парки Москвы» и ГБУ «Мосприрода»;",
        "въезд, передвижение транспортных средств по ООЗТ вне установленных маршрутов, а так же остановка и стоянка транспортных средств вне установленных мест будет запрещена;",
        "Проезд будет осуществляться по существующей дорожной сети с исключением заезда автотранспорта на территории, занятые зелеными насаждениями.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"- {bullet}")
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Также будет осуществлен видеоконтроль автомобильных номеров въезжающего "
        "на территорию ландшафтного заказника «Теплый Стан» автотранспорта по адресу: "
        "г. Москва, ул. Островитянова, вл.10."
    )
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Председатель")
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run("Правления АНО СК «ОЛИМПИК»")
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("А.П. Дмитриенко")
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Исполнитель")
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run("Пруцков Илья Евгеньевич")
    run.font.size = Pt(12)
    p = doc.add_paragraph()
    run = p.add_run("+79166571350")
    run.font.size = Pt(12)

    date_fn = target_date.strftime("%d_%m_%Y")
    filename = f"Razovye_propuska_{date_fn}.docx"
    filepath = Path("/tmp") / filename
    doc.save(str(filepath))
    logger.info(f"Документ (письмо) сохранён: {filepath}")
    return filepath


def generate_car_list(cars, target_date_str):
    """Генерирует Word-документ с простым перечнем автомобилей."""
    target_date = datetime.strptime(target_date_str, "%d.%m.%Y")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Перечень автомобилей на {target_date_str}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Таблица
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["№", "Марка автомобиля", "Номер автомобиля"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        set_cell_border(cell)

    for idx, car in enumerate(cars, 1):
        row = table.add_row()
        for i, val in enumerate([str(idx), car["car_brand"], car["car_number"]]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            set_cell_border(cell)

    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(7)

    date_fn = target_date.strftime("%d_%m_%Y")
    filename = f"Perechen_auto_{date_fn}.docx"
    filepath = Path("/tmp") / filename
    doc.save(str(filepath))
    logger.info(f"Документ (перечень) сохранён: {filepath}")
    return filepath


def send_telegram_document(filepath, caption):
    """Отправляет файл через Telegram Bot API."""
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendDocument"
    boundary = "----FormBoundary7MA4YWxkTrZu0gW"
    filename = filepath.name
    with open(filepath, "rb") as f:
        file_data = f.read()
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="chat_id"\r\n\r\n'
        f"{RECIPIENT_CHAT_ID}\r\n"
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="caption"\r\n\r\n'
        f"{caption}\r\n"
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="document"; filename="{filename}"\r\n'
        f"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
    ).encode("utf-8") + file_data + f"\r\n--{boundary}--\r\n".encode("utf-8")
    req = urllib.request.Request(
        url, data=body,
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
        method="POST",
    )
    resp = urllib.request.urlopen(req)
    result = json.loads(resp.read())
    if result.get("ok"):
        logger.info(f"Файл отправлен в Telegram: {filename}")
    else:
        logger.error(f"Ошибка отправки: {result}")
        sys.exit(1)


def send_telegram_message(text):
    """Отправляет текстовое сообщение через Telegram Bot API."""
    url = f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage"
    data = json.dumps({"chat_id": RECIPIENT_CHAT_ID, "text": text}).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"}, method="POST")
    resp = urllib.request.urlopen(req)
    result = json.loads(resp.read())
    if result.get("ok"):
        logger.info("Сообщение отправлено в Telegram")
    else:
        logger.error(f"Ошибка отправки: {result}")


def main():
    if not BOT_TOKEN:
        logger.error("BOT_TOKEN не задан")
        sys.exit(1)

    now = datetime.now(TIMEZONE)
    tomorrow = now + timedelta(days=1)
    target_date_str = tomorrow.strftime("%d.%m.%Y")

    logger.info(f"Генерация пропусков на {target_date_str}")
    cars = get_cars_for_date(target_date_str)

    if not cars:
        msg = f"На {target_date_str} заявок на проезд нет."
        logger.info(msg)
        send_telegram_message(msg)
        return

    # Документ 1: Официальное письмо
    letter_path = generate_document(cars, target_date_str)
    send_telegram_document(letter_path, f"📄 Письмо на разовый проезд на {target_date_str}\nАвтомобилей: {len(cars)}")
    letter_path.unlink(missing_ok=True)

    # Документ 2: Перечень автомобилей
    list_path = generate_car_list(cars, target_date_str)
    send_telegram_document(list_path, f"📋 Перечень автомобилей на {target_date_str}")
    list_path.unlink(missing_ok=True)

    logger.info("Готово! Оба документа отправлены.")


if __name__ == "__main__":
    main()
